import os, subprocess
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"C:\Users\Dee\Desktop\mobicova-platform"
PANDOC = r"C:\Users\Dee\AppData\Local\Pandoc\pandoc.exe"
LOGO = os.path.join(ROOT, "client", "public", "images", "logo.png")
BANNER = os.path.join(ROOT, "_banner.png")
TEAL = RGBColor(0x0A, 0x7B, 0x7B)
MD = "MobiCova-USSD-Menu-Walkthrough.md"
DOCX = "MobiCova-USSD-Menu-Walkthrough.docx"

# header banner: white logo on teal gradient + orange rule
W, H = 2400, 470
top, bot = (12, 122, 122), (8, 52, 52)
grad = Image.new("RGBA", (1, H))
for y in range(H):
    t = y / H
    grad.putpixel((0, y), tuple(int(top[i] + (bot[i] - top[i]) * t) for i in range(3)) + (255,))
banner = grad.resize((W, H)).convert("RGBA")
ImageDraw.Draw(banner).rectangle([0, H - 16, W, H], fill=(208, 126, 19, 255))
logo = Image.open(LOGO).convert("RGBA")
th = 250
logo = logo.resize((int(logo.width * th / logo.height), th), Image.LANCZOS)
banner.alpha_composite(logo, (100, (H - 16 - th) // 2))
banner.convert("RGB").save(BANNER)

subprocess.run([PANDOC, MD, "-o", DOCX, "--toc", "--toc-depth=2"], check=True, cwd=ROOT)

doc = Document(os.path.join(ROOT, DOCX))
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)
for name, size in [("Title", 24), ("Heading 1", 17), ("Heading 2", 13.5), ("Heading 3", 12)]:
    try:
        st = doc.styles[name]
        st.font.name = "Calibri"; st.font.bold = True; st.font.size = Pt(size); st.font.color.rgb = TEAL
    except KeyError:
        pass
p = doc.paragraphs[0].insert_paragraph_before()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(BANNER, width=Inches(6.3))
doc.save(os.path.join(ROOT, DOCX))

os.remove(BANNER)
print("branded:", DOCX)
